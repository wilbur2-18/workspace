from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mac/Desktop/workspace/1 综合分析平台")
OUT = ROOT / "0 产品文档/综合分析平台v2.0（用户操作手册）20260527.docx"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 45, 58)
MUTED = RGBColor(99, 110, 123)
BORDER = "B7C7D9"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
NOTE_FILL = "FFF8E6"
TABLE_WIDTH_DXA = 9360


SCREENSHOTS = [
    ("截图-01", "工作台首页", "顶部导航、工作台卡片、搜索与筛选入口", "生产系统工作台首页整屏"),
    ("截图-02", "创建工作台", "工作台名称、简介、权限范围", "创建工作台弹窗"),
    ("截图-03", "工作台筛选", "我的工作台/共享工作台、搜索、排序", "工作台列表筛选状态"),
    ("截图-04", "技能库首页", "我的技能/共享技能、创建技能、筛选、技能卡片", "技能库首页整屏"),
    ("截图-05", "创建技能", "技能名称、描述、标签、继续配置", "创建技能弹窗"),
    ("截图-06", "技能配置", "基本信息、审计思路、文件树、保存", "技能配置编辑界面"),
    ("截图-07", "历史版本", "版本列表、版本详情、回退入口", "技能历史版本页签"),
    ("截图-08", "审计助手总览", "资源/技能左栏、对话区、任务/结果右栏", "审计助手工作台整屏"),
    ("截图-09", "资源区", "文件、库表、图谱、知识库入口与搜索", "审计助手资源页签"),
    ("截图-10", "上传资料", "上传入口、文件列表、开始上传", "上传资料弹窗"),
    ("截图-11", "上传解析状态", "上传中、排队中、解析中、失败、完成", "资源列表状态筛选或状态行"),
    ("截图-12", "添加库表", "数据源、表格勾选、搜索、添加到工作台", "添加库表弹窗"),
    ("截图-13", "添加图谱", "图谱卡片、详情、勾选添加", "添加数据图谱弹窗"),
    ("截图-14", "工作台技能", "引用技能、创建技能、技能列表、匹配状态", "审计助手技能页签"),
    ("截图-15", "对话引用", "@ 引用、/ 技能、输入框、发送按钮", "审计助手对话输入区"),
    ("截图-16", "创建单次任务", "选择技能、任务要求、输出位置", "创建任务第 1 步"),
    ("截图-17", "选择任务资源", "文件/结果/库表/图谱/知识库与已选资源", "创建任务资源选择步骤"),
    ("截图-18", "创建跑批任务", "数据源文件、标识列、任务指令", "跑批任务配置弹窗"),
    ("截图-19", "任务列表", "任务状态、筛选、详情、重跑、中止、删除", "右侧任务页签"),
    ("截图-20", "结果树", "结果文件夹、排序、搜索、批量打包入口", "右侧结果页签"),
    ("截图-21", "结果预览", "基本信息、输出结果、编辑、下载、历史版本", "结果预览或内嵌详情"),
    ("截图-22", "资源预览", "文件预览、OCR 结果、引用来源", "资料预览弹窗"),
    ("截图-23", "系统管理入口", "系统管理左侧菜单和审计人员可见入口", "系统管理首页"),
    ("截图-24", "数据源管理", "数据源列表、新建/编辑、测试连接", "数据源管理页面或弹窗"),
]

FIGURE_COUNTER = {}


def set_run_font(run, size=None, color=None, bold=None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_para_font(paragraph, size=11, color=INK, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold if bold else None)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color="D9E1EA", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "3")
        node.set(qn("w:color"), color)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, width_dxa=TABLE_WIDTH_DXA):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def style_table(table, header=True, widths=None):
    table.autofit = False
    set_table_fixed_width(table)
    if widths:
        set_table_width(table, widths)
    for r_idx, row in enumerate(table.rows):
        if header and r_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                tbl_header = OxmlElement("w:tblHeader")
                tr_pr.append(tbl_header)
            tbl_header.set(qn("w:val"), "true")
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if header and r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    set_para_font(paragraph, size=10.5, color=INK, bold=True)
            else:
                for paragraph in cell.paragraphs:
                    set_para_font(paragraph, size=10.5, color=INK)


def add_footer(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("综合分析平台 v2.0 用户操作手册 ｜ 第 ")
    set_run_font(run, size=9, color=MUTED)
    add_field(paragraph, "PAGE", display_text="1", size=9, color=MUTED)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_field(paragraph, instruction, display_text="1", size=10.5, color=INK):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    text = OxmlElement("w:t")
    text.text = display_text
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    set_run_font(run, size=size, color=color)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)
        add_footer(sec)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name, size, color, bold, after in [
        ("Manual Cover Title", 28, RGBColor(0, 0, 0), True, 10),
        ("Manual Cover Subtitle", 22, DARK_BLUE, True, 8),
        ("Manual Meta", 11, MUTED, False, 6),
        ("Manual Label", 11, DARK_BLUE, True, 4),
        ("Manual Screenshot", 10, DARK_BLUE, False, 8),
        ("Manual Caption", 9, MUTED, False, 8),
    ]:
        if style_name in styles:
            style = styles[style_name]
        else:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title_page(doc):
    p = doc.add_paragraph(style="Manual Cover Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(10)
    p.add_run("综合分析平台 v2.0")

    p = doc.add_paragraph(style="Manual Cover Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("用户操作手册")

    p = doc.add_paragraph(style="Manual Meta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)

    data = [
        ("文档版本", "v1.0"),
        ("发布日期", "2026年5月27日"),
    ]
    for k, v in data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{k}：")
        set_run_font(r, size=10.5, color=INK, bold=True)
        r = p.add_run(v)
        set_run_font(r, size=10.5, color=INK)

    doc.add_page_break()


def add_document_notes_page(doc):
    doc.add_heading("文档说明", level=1)
    doc.add_paragraph(
        "本手册用于指导一线审计人员使用综合分析平台 v2.0，围绕工作台完成资源组织、技能调用、审计助手协同、任务执行和结果沉淀。"
    )

    rows = [
        ("适用范围", "工作台、技能、资源、对话、结果、任务、系统管理入口及常见状态处理。管理员配置仅说明入口，不作为管理员手册展开。"),
        ("读者对象", "一线审计人员、审计项目组成员、培训交付人员，以及需要理解业务使用路径的项目交付人员。"),
        ("文档口径", "按审计任务组织操作路径，帮助审计人员在工作台内完成资料组织、技能调用、对话分析、任务执行和结果沉淀。"),
        ("权限说明", "审计人员实际可见的工作台、技能、资源、对话、结果、任务和系统管理入口，以实际部署和权限配置为准。"),
        ("脱敏说明", "文中示例均采用业务可读脱敏口径。制作界面截图时，应遮挡真实姓名、单位、项目、金额、账号、内网地址等敏感信息。"),
    ]
    for label, body in rows:
        add_labeled_paragraph(doc, label, body)

    doc.add_page_break()
    doc.add_heading("核心概念与相关概念", level=1)
    terms = [
        ("工作台", "围绕一次审计事项组织技能、资源、对话、任务和结果的业务空间。审计人员通常先进入或创建工作台，再在该工作台内开展分析。我的工作台用于个人负责事项，共享工作台用于协同审计，审计助手是进入工作台后的主操作界面。"),
        ("技能", "可复用的审计分析方法，用于描述分析思路、资料要求和输出要求，可在对话或任务中调用。技能库用于统一查找和维护技能；共享技能用于复用成熟方法；工作台技能是当前工作台已引用或创建、可直接使用的技能。"),
        ("资源", "审计分析的上下文来源，包括资料文件、库表、图谱、知识库入口和已有结果。资源进入工作台后，才能被对话或任务引用。资料文件是合同、付款资料、会议纪要等文件类材料；资源树用于按业务类别组织资源。"),
        ("对话", "审计人员在审计助手中提出分析问题、引用资源、调用技能并接收系统回复的交互过程。@ 引用用于指定资料、结果、库表或图谱；/ 技能用于调用分析方法；有复用价值的对话产出可保存为结果或沉淀为技能。"),
        ("结果", "对话保存、任务执行或人工编辑形成的分析产物，供审计人员复核、引用、下载或继续加工。结果树用于组织分析产物；结果预览用于查看内容和引用来源；历史版本用于追溯修改；打包下载用于批量导出结果。"),
        ("任务", "由技能、资源和执行要求组成的分析执行单元，用于承载可跟踪、可重跑的分析过程。单次任务用于一次性分析；跑批任务用于对多行对象或多个主体重复执行；子任务用于查看跑批明细；任务状态用于判断排队、执行、完成或失败。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "术语"
    table.rows[0].cells[1].text = "概念说明"
    for term, desc in terms:
        cells = table.add_row().cells
        cells[0].text = term
        cells[1].text = desc
    style_table(table, widths=[0.65, 5.65])

    doc.add_page_break()


def add_manual_toc(doc):
    doc.add_heading("目录", level=1)
    items = [
        ("核心概念与相关概念", "3"),
        ("1 使用前说明", "6"),
        ("2 快速上手：完成一次审计分析闭环", "6"),
        ("3 工作台管理", "8"),
        ("4 资源与资料文件管理", "11"),
        ("5 技能使用与技能库", "15"),
        ("6 审计助手协同分析", "18"),
        ("7 任务管理", "21"),
        ("8 结果管理", "24"),
        ("9 典型审计场景", "27"),
        ("10 系统管理入口", "28"),
        ("11 FAQ 与状态说明", "30"),
        ("附录 待确认事项", "31"),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(title)
        p.add_run("\t{}".format(page))
        set_para_font(p, size=10.5, color=INK)
    doc.add_page_break()


def add_callout(doc, title, body, fill=LIGHT_FILL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color="D9E1EA")
    r = p.add_run(f"{title}：")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=INK)


def add_labeled_paragraph(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label}：")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=INK)


def add_section_label(doc, label):
    p = doc.add_paragraph(style="Manual Label")
    p.add_run(label)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.22)
        p.add_run(item)
        set_para_font(p)


def add_placeholder(doc, shot_id):
    shot = next(item for item in SCREENSHOTS if item[0] == shot_id)
    chapter = getattr(add_placeholder, "chapter", "0")
    FIGURE_COUNTER[chapter] = FIGURE_COUNTER.get(chapter, 0) + 1
    figure_id = "截图 {}-{}".format(chapter, FIGURE_COUNTER[chapter])
    p = doc.add_paragraph(style="Manual Screenshot")
    p.add_run("{}：{}".format(figure_id, shot[1]))
    p = doc.add_paragraph(style="Manual Caption")
    p.add_run("建议截取区域：{}；需突出内容：{}。".format(shot[3], shot[2]))


def add_steps(doc, steps):
    for idx, step in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run("{}. {}".format(idx, step))
        set_para_font(p)


def add_bullets(doc, bullets):
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)
        set_para_font(p)


def get_start_checks(title, prereq):
    checks_by_title = {
        "2.1 从工作台进入审计助手": ["已登录系统。", "能在「我的工作台」或「共享工作台」中看到目标工作台。"],
        "2.2 完成一次资料到结果的分析": ["已进入目标工作台。", "左侧资源区已有本次要分析的资料或数据。", "工作台内已有可用技能；如没有技能，先参考第 5 章准备。"],
        "3.1 创建工作台": ["能看到顶部导航中的「工作台」。", "页面上显示「创建工作台」入口。"],
        "3.2 查找和进入工作台": ["目标工作台已经创建。", "能在「我的工作台」或「共享工作台」中看到相关列表。"],
        "3.3 编辑或删除工作台": ["能在工作台列表中看到目标工作台。", "工作台卡片上显示编辑、删除或更多操作入口。", "删除前已确认该工作台中的资料、任务和结果不需要继续保留。"],
        "4.1 上传文件资料": ["已进入目标工作台的审计助手。", "待上传资料已经按单位要求完成脱敏或合规确认。", "能看到左侧「资源」页签和上传入口。"],
        "4.2 预览资料和 OCR 结果": ["目标资料已经上传到资源区。", "资料状态允许预览或查看识别结果。"],
        "4.3 新建文件夹并整理资料": ["已进入目标工作台的资源区。", "资源区中已有需要整理的资料，或准备在文件夹下继续上传资料。"],
        "4.4 添加数据库表": ["资源区能看到数据库表入口。", "目标数据源和表在当前账号下可见。"],
        "4.5 添加数据图谱": ["资源区能看到数据图谱入口。", "目标图谱在当前账号下可见。"],
        "5.1 查找并复用技能库技能": ["能进入「技能库」。", "需要复用的技能在「我的技能」或「共享技能」中可见。"],
        "5.2 在工作台内引用技能": ["已进入目标工作台。", "左侧「技能」页签可见。", "目标技能在技能库中可见。"],
        "5.3 创建和配置技能": ["能看到「创建技能」入口。", "已明确技能名称、适用资料和分析思路。"],
        "5.4 管理技能版本和共享": ["目标技能已创建。", "技能详情或配置页面中能看到历史版本、共享或复制入口。"],
        "6.1 使用对话引用资料和技能": ["已进入目标工作台的审计助手。", "需要引用的资源或技能已经加入当前工作台。"],
        "6.2 将对话产出保存为结果": ["对话中已有需要保留的分析内容。", "助手回答区域显示保存结果或保存到结果树的入口。"],
        "6.3 将对话经验总结为技能": ["当前对话已经形成稳定的资料范围、判断逻辑和输出口径。", "页面中显示总结为技能或生成技能入口。"],
        "7.1 创建单次任务": ["当前工作台已有可用技能。", "需要参与分析的资源已加入工作台。", "能看到创建任务入口。"],
        "7.2 创建跑批任务": ["已有可用技能。", "已准备跑批数据源文件或结构化数据。", "能看到跑批任务入口。"],
        "7.3 查看、重跑、中止和删除任务": ["右侧任务页签中已有任务记录。", "需要处理的任务显示详情、重跑、中止或删除入口。"],
        "7.4 处理跑批子任务": ["已创建跑批任务。", "跑批父任务可打开子任务明细。"],
        "8.1 查看和组织结果树": ["右侧「结果」页签可见。", "当前工作台已有结果或结果文件夹。"],
        "8.2 编辑、保存和下载结果": ["目标结果已经生成。", "结果预览页显示编辑、保存、历史版本或下载入口。"],
        "8.3 打包下载结果": ["结果树中已有需要导出的结果或文件夹。", "页面显示打包下载入口。"],
        "10.1 查看系统管理入口": ["顶部导航中能看到「系统管理」。"],
        "10.2 了解数据源管理入口": ["能进入「系统管理」。", "左侧菜单中能看到数据源管理入口。"],
    }
    return checks_by_title.get(title, [prereq])


def add_operation(doc, title, scenario, prereq, entrance, steps, result, notes, shot_ids=None):
    doc.add_heading(title, level=2)
    doc.add_paragraph(scenario)
    add_labeled_paragraph(doc, "入口", entrance)
    add_section_label(doc, "开始前确认")
    add_checklist(doc, get_start_checks(title, prereq))
    add_section_label(doc, "操作步骤")
    add_steps(doc, steps)
    add_labeled_paragraph(doc, "完成后", result)

    if notes:
        add_callout(doc, "注意事项", "；".join(notes), fill=NOTE_FILL)
    if shot_ids:
        add_placeholder.chapter = title.split(".", 1)[0]
        for shot_id in shot_ids:
            add_placeholder(doc, shot_id)


def add_section_1(doc):
    doc.add_heading("1 使用前说明", level=1)
    doc.add_paragraph(
        "本章只保留使用系统前必须明确的人工复核原则。适用对象、权限和脱敏口径已在文档说明中统一说明，正文不再重复展开。"
    )

    doc.add_heading("1.1 人工复核原则", level=2)
    doc.add_paragraph(
        "系统生成的分析结果用于辅助审计判断。审计人员应结合原始资料、审计程序和专业判断进行复核确认，不应将系统生成内容直接作为审计结论。"
    )


def add_section_2(doc):
    doc.add_heading("2 快速上手：完成一次审计分析闭环", level=1)
    doc.add_paragraph("本章提供一条最短可用路径，适合新审计人员第一次使用系统时照着完成。熟练人员可直接查阅后续分章。")
    add_operation(
        doc,
        "2.1 从工作台进入审计助手",
        "需要围绕一个审计事项开展资料整理和分析。",
        "已登录系统，且当前账号具备查看或创建工作台的权限。",
        "在顶部导航进入「工作台」，从工作台列表打开目标工作台。",
        [
            "进入系统后，打开「工作台」。",
            "在「我的工作台」或「共享工作台」中搜索目标工作台。",
            "点击工作台卡片或「继续审计」，进入该工作台的审计助手。",
        ],
        "系统打开对应工作台的审计助手，页面展示资源区、对话区、任务/结果区。",
        ["如果看不到目标工作台，先确认是否切换到共享工作台，或联系管理员检查授权范围。"],
        ["截图-01", "截图-08"],
    )
    add_operation(
        doc,
        "2.2 完成一次资料到结果的分析",
        "已进入工作台，需要基于资料和技能生成可复核的分析结果。",
        "工作台中已有资料和可用技能；如没有资料或技能，先参考第 4、5 章完成准备。",
        "在目标工作台的「审计助手」页面，从左侧资源区、技能区和中间对话区开始操作。",
        [
            "在左侧资源区确认本次要分析的资料、库表或图谱已加入工作台。",
            "在对话输入框中使用 @ 引用资料或结果，必要时使用 / 选择技能。",
            "发送分析要求，等待审计助手生成分析过程或建议创建任务。",
            "如需要创建任务，按提示选择技能。",
            "选择参与分析的资源。",
            "确认输出位置并提交任务。",
            "在右侧任务区查看执行状态。",
            "任务完成后，在结果树中打开生成结果。",
            "复核结果内容，必要时编辑、保存或下载。"
        ],
        "任务状态为完成，结果树出现对应结果，审计人员可预览、编辑、下载或继续对话追问。",
        ["系统生成结果需由审计人员复核确认后再作为正式结论依据。"],
        ["截图-15", "截图-16", "截图-19", "截图-21"],
    )


def add_section_3(doc):
    doc.add_heading("3 工作台管理", level=1)
    doc.add_paragraph("工作台是综合分析平台的业务容器。审计人员应先确认分析对象属于哪个工作台，再在该工作台内组织资料、技能、任务和结果。")
    add_operation(
        doc,
        "3.1 创建工作台",
        "需要为新的审计事项建立独立分析空间。",
        "当前账号具备创建工作台权限。",
        "在顶部导航进入「工作台」，点击「创建工作台」。",
        [
            "在顶部导航进入「工作台」。",
            "点击「创建工作台」。",
            "填写工作台名称和简介，名称建议包含审计事项、年度或单位范围。",
            "选择权限范围。可选择仅自己可见，或按实际配置选择指定部门、指定用户可见。",
            "点击「创建」。"
        ],
        "系统创建工作台，并进入该工作台的审计助手或在工作台列表显示新工作台。",
        ["权限范围会影响其他审计人员是否可见；真实权限以生产部署配置为准。"],
        ["截图-02"],
    )
    add_operation(
        doc,
        "3.2 查找和进入工作台",
        "工作台较多时，需要快速找到目标工作台。",
        "目标工作台已创建，当前账号有查看权限。",
        "在顶部导航进入「工作台」，使用列表上方搜索、筛选和排序入口。",
        [
            "在「我的工作台」和「共享工作台」之间切换，确认工作台来源。",
            "在搜索框输入工作台名称、审计事项或关键字。",
            "按创建时间、更新时间或系统提供的排序方式筛选。",
            "点击工作台卡片或「继续审计」进入。"
        ],
        "进入目标工作台后，审计助手会加载该工作台下的资料、技能、任务和结果。",
        ["搜索不到不代表工作台不存在，可能是权限、关键字或筛选条件导致。"],
        ["截图-03"],
    )
    add_operation(
        doc,
        "3.3 编辑或删除工作台",
        "工作台名称、简介需要规范化，或误建工作台需要清理。",
        "当前账号具备该工作台的编辑或删除权限；删除前已确认无需保留其中资料、任务和结果。",
        "在「工作台」列表中找到目标工作台，打开工作台卡片的更多操作入口。",
        [
            "在工作台列表中找到目标工作台。",
            "打开工作台卡片上的更多操作。",
            "选择编辑时，修改工作台名称或简介并保存。",
            "选择删除时，先阅读系统提示，确认是否会影响资料、任务和结果。",
            "确认无误后再提交操作。"
        ],
        "编辑后，工作台列表和审计助手标题同步更新；删除后，工作台从当前列表移除。",
        ["删除不是归档手段。正式生产系统如有归档、停用或移交能力，应优先按单位管理规则处理。"],
        ["截图-02", "截图-03"],
    )


def add_section_4(doc):
    doc.add_heading("4 资源与资料文件管理", level=1)
    doc.add_paragraph("资源是分析的上下文来源。审计人员可在工作台中管理资料文件、数据库表、数据图谱和知识库入口，并将这些资源引用到对话或任务中。")
    add_operation(
        doc,
        "4.1 上传文件资料",
        "需要将审计材料加入当前工作台。",
        "已进入目标工作台；已准备脱敏或合规可上传的资料文件。",
        "在审计助手左侧「资源」页签的文件资源区域，点击上传入口。",
        [
            "在审计助手左侧选择「资源」。",
            "进入文件资源区域，点击上传入口。",
            "在上传弹窗中选择文件，或将文件拖入上传区域。",
            "确认文件列表无误后，点击「开始上传」。",
            "返回资源列表，查看上传和解析状态。"
        ],
        "文件出现在资源列表中，并进入上传中、排队中、解析中、完成或失败等状态。",
        ["真实可上传格式、大小和数量以生产系统配置为准；上传后请等待解析完成再用于正式分析。"],
        ["截图-09", "截图-10", "截图-11"],
    )
    add_operation(
        doc,
        "4.2 预览资料和 OCR 结果",
        "需要确认资料内容、识别结果或引用来源是否正确。",
        "资料已上传并具备可预览状态。",
        "在审计助手左侧「资源」页签中，点击目标资料名称或预览入口。",
        [
            "在资源列表中选择目标资料。",
            "点击预览或资料名称，打开资料预览。",
            "查看文件基本信息、原始预览和 OCR/识别结果。",
            "如页面提供普通视图和行号视图，可按复核需要切换。",
            "确认资料可用于分析后，返回审计助手。"
        ],
        "资料预览页显示基本信息和识别结果，审计人员确认资料可被引用到对话或任务中。",
        ["OCR 或解析结果不完整时，应先补充资料或重新处理，不建议直接作为最终结论依据。"],
        ["截图-22"],
    )
    add_operation(
        doc,
        "4.3 新建文件夹并整理资料",
        "工作台资料较多，需要按合同、付款、验收、会议纪要等业务类别归档。",
        "已进入目标工作台；当前账号具备资料管理权限。",
        "在审计助手左侧「资源」页签的文件资源区域，使用新建文件夹或移动入口。",
        [
            "在资源区选择文件资源。",
            "点击新建文件夹入口，输入业务可读的文件夹名称。",
            "将已上传资料移动到对应文件夹，或在文件夹下继续上传资料。",
            "需要调整时，可对文件夹执行重命名、删除或移动。",
            "整理完成后，在资源树中确认资料层级是否清晰。"
        ],
        "资料按业务类别在资源树中展示，后续可按资料或文件夹引用到对话和任务。",
        ["文件夹名称应表达业务含义，不建议使用个人姓名、临时编号或未经脱敏的内部编号。"],
        ["截图-09"],
    )
    add_operation(
        doc,
        "4.4 添加数据库表",
        "需要将结构化数据表加入当前工作台，用于查询、对比或分析。",
        "生产环境已配置可用数据源；当前账号具备访问该数据源的权限。",
        "在审计助手左侧「资源」页签中，进入数据库表资源入口并点击「添加库表」。",
        [
            "在资源区选择数据库表入口。",
            "点击「添加库表」。",
            "选择数据源，并按表名或注释搜索目标表。",
            "勾选需要加入工作台的表。",
            "点击添加，将库表加入当前工作台。"
        ],
        "所选库表显示在工作台资源区，可预览字段或引用到对话中。",
        ["库表可见范围和字段权限以生产数据源授权为准。"],
        ["截图-12"],
    )
    add_operation(
        doc,
        "4.5 添加数据图谱",
        "需要将图谱资源作为审计分析上下文。",
        "生产环境已配置数据图谱资源，且当前账号有访问权限。",
        "在审计助手左侧「资源」页签中，进入数据图谱资源入口并点击添加图谱。",
        [
            "在资源区选择数据图谱入口。",
            "点击添加图谱。",
            "查看图谱卡片，确认图谱名称、实体和关系信息。",
            "勾选需要加入工作台的图谱。",
            "点击添加。"
        ],
        "图谱出现在工作台资源区，可在对话或任务中引用。",
        ["图谱内容和权限以生产系统配置为准；知识库入口如未开放，仅作为可见入口说明。"],
        ["截图-13"],
    )


def add_section_5(doc):
    doc.add_heading("5 技能使用与技能库", level=1)
    doc.add_paragraph("技能用于沉淀可复用的分析方法。审计人员通常有两类场景：在工作台内引用技能完成任务，以及在技能库中创建、维护或复用技能。")
    add_operation(
        doc,
        "5.1 查找并复用技能库技能",
        "需要复用已有分析方法，例如金额核对、进度偏差扫描、供应商交叉比对。",
        "已进入技能库；技能库中已有我的技能或共享技能。",
        "在顶部导航进入「技能库」，使用「我的技能」或「共享技能」列表。",
        [
            "进入「技能库」。",
            "在「我的技能」和「共享技能」之间切换。",
            "输入技能名称、业务场景或标签关键词搜索。",
            "按标签、更新时间或系统提供的排序方式缩小范围。",
            "打开技能详情，确认适用场景、审计思路和历史版本。",
            "如为共享技能且需要修改，先添加或复制到我的技能。"
        ],
        "技能详情页可打开；目标技能可被添加到我的技能或引用到工作台。",
        ["共享技能通常用于复用成熟方法；如需修改，应先形成个人或工作台副本，避免影响原技能。"],
        ["截图-04"],
    )
    add_operation(
        doc,
        "5.2 在工作台内引用技能",
        "当前工作台需要使用已有技能进行分析。",
        "技能库中已有可用技能，当前账号具备查看或引用权限。",
        "在审计助手左侧切换到「技能」页签，点击添加或引用技能入口。",
        [
            "在审计助手左侧切换到「技能」。",
            "点击添加或引用技能入口。",
            "在弹窗中切换「我的技能」或「共享技能」。",
            "通过搜索、标签或筛选找到目标技能。",
            "勾选技能并点击「引用到当前工作台」。"
        ],
        "技能出现在当前工作台技能列表中，可用于对话引用或创建任务。",
        ["技能库中的技能和工作台内已引用技能是两个使用场景，手册中分别说明，避免混写。"],
        ["截图-14"],
    )
    add_operation(
        doc,
        "5.3 创建和配置技能",
        "需要把常用审计分析方法沉淀为可复用技能。",
        "当前账号具备创建技能权限；已明确技能名称、适用场景和分析思路。",
        "在「技能库」点击「创建技能」，或在工作台技能区点击创建技能入口。",
        [
            "进入「技能库」或在工作台技能区点击创建技能。",
            "填写技能名称、说明和标签。",
            "进入技能配置界面，补充审计思路和必要的技能文件。",
            "检查文件树、配置内容和保存状态。",
            "点击保存。"
        ],
        "技能保存成功后，可在我的技能或当前工作台技能列表中查看和使用。",
        ["技能说明应写清适用资料、判断逻辑和输出要求，避免只写宽泛目标。"],
        ["截图-04", "截图-05", "截图-06"],
    )
    add_operation(
        doc,
        "5.4 管理技能版本和共享",
        "需要保留技能修改记录，或将成熟技能复用给其他审计人员。",
        "技能已创建；当前账号具备编辑、共享或版本操作权限。",
        "在「技能库」打开目标技能详情或配置页面，进入历史版本或共享入口。",
        [
            "打开目标技能的配置或详情。",
            "在历史版本中查看版本记录和变更说明。",
            "需要回退时，选择目标版本并按系统提示确认。",
            "需要共享时，使用共享入口将技能发布到共享范围。",
            "其他审计人员可在共享技能中复制或添加到自己的技能库。"
        ],
        "历史版本列表显示目标版本；共享状态在技能卡片或详情页更新。",
        ["共享范围、审批和发布规则以生产环境配置为准。"],
        ["截图-07"],
    )


def add_section_6(doc):
    doc.add_heading("6 审计助手协同分析", level=1)
    doc.add_paragraph("审计助手是工作台内的主分析界面。审计人员在此引用资料、调用技能、发起任务、查看工具执行过程，并将结论保存为结果。")
    add_operation(
        doc,
        "6.1 使用对话引用资料和技能",
        "需要围绕资料、结果或结构化资源进行追问、总结或比对。",
        "已进入工作台；资源或技能已加入当前工作台。",
        "在工作台的「审计助手」页面，使用中间对话输入区。",
        [
            "在对话输入框输入分析问题。",
            "使用 @ 引用资料、结果、库表或图谱。",
            "使用 / 选择需要调用的技能。",
            "检查引用对象是否正确。",
            "点击发送。"
        ],
        "审计助手生成回答、工具执行过程或任务建议，引用标记可用于追溯上下文。",
        ["引用对象越准确，后续结果越容易复核；不要只输入泛泛问题而不指定资料范围。"],
        ["截图-15"],
    )
    add_operation(
        doc,
        "6.2 将对话产出保存为结果",
        "对话中形成了需要沉淀和复核的结论。",
        "对话已有可保存内容；当前账号具备保存结果权限。",
        "在审计助手对话消息或助手回答区域，使用保存到结果树的入口。",
        [
            "在对话消息或助手回答中找到保存入口。",
            "选择保存为结果或保存到结果树。",
            "填写结果名称，必要时选择文件夹位置。",
            "确认保存。",
            "在右侧结果树中打开保存后的结果。"
        ],
        "对话产出进入结果树，后续可编辑、下载、引用或纳入任务上下文。",
        ["保存结果不等于最终定稿，正式使用前仍需人工复核。"],
        ["截图-20", "截图-21"],
    )
    add_operation(
        doc,
        "6.3 将对话经验总结为技能",
        "多轮对话中形成了可重复使用的核查方法，希望沉淀为工作台技能或后续技能库资产。",
        "当前对话已有明确的资料范围、判断逻辑和输出口径；审计人员确认该方法具备复用价值。",
        "在审计助手对话回复下方，或左侧「技能」页签中使用生成技能入口。",
        [
            "在对话回复下方选择总结为技能，或从技能区选择生成技能入口。",
            "填写技能名称和生成要求。",
            "检查系统生成的审计资料要求、审计思路和输出要求。",
            "必要时手工修改技能配置。",
            "保存为工作台级技能。",
            "回到左侧技能页签，确认新技能已出现。"
        ],
        "对话经验沉淀为工作台技能，可继续用于任务、对话引用或后续入库。",
        ["只有稳定、可复核、可重复执行的方法才适合沉淀为技能；一次性判断和未经验证的结论不建议入库。"],
        ["截图-14", "截图-15"],
    )


def add_section_7(doc):
    doc.add_heading("7 任务管理", level=1)
    doc.add_paragraph("任务用于承载可跟踪的分析执行。审计人员应通过任务状态判断执行进度，并在失败或结果不符合预期时重跑或调整配置。")
    add_operation(
        doc,
        "7.1 创建单次任务",
        "需要针对一组已选资料执行一次分析。",
        "当前工作台已有可用技能和资源。",
        "在审计助手中点击创建任务入口，选择单次任务。",
        [
            "在审计助手中选择创建任务。",
            "选择任务类型为单次任务。",
            "选择要执行的技能。",
            "填写或确认任务要求和输出位置。",
            "选择需要参与分析的资源。",
            "提交任务。"
        ],
        "任务进入任务列表，状态从排队或执行中逐步更新，完成后生成结果。",
        ["任务要求应包含分析目标、关注字段和输出格式，避免只写“帮我分析”。"],
        ["截图-16", "截图-17"],
    )
    add_operation(
        doc,
        "7.2 创建跑批任务",
        "需要对多行对象、多个企业或多组资料重复执行同类分析。",
        "已准备符合生产系统要求的数据源文件或结构化数据；已有可用技能。",
        "在审计助手中点击创建任务入口，选择跑批任务。",
        [
            "在创建任务时选择跑批任务。",
            "上传或选择跑批数据源。",
            "选择用于识别每一行对象的标识列。",
            "确认系统生成或审计人员编辑的任务指令。",
            "选择执行资源和输出位置。",
            "提交跑批任务。"
        ],
        "系统生成跑批父任务及子任务，审计人员可查看整体状态和子任务明细。",
        ["跑批数据源格式、列识别规则和并发策略以生产系统实际配置为准。"],
        ["截图-18"],
    )
    add_operation(
        doc,
        "7.3 查看、重跑、中止和删除任务",
        "需要跟踪任务执行状态或处理失败任务。",
        "任务已创建并出现在任务列表中。",
        "在审计助手右侧切换到「任务」页签，打开任务列表或任务详情。",
        [
            "在右侧切换到任务页签。",
            "按状态筛选排队中、执行中、完成或失败任务。",
            "点击任务查看详情、引用资源、输出位置和执行过程。",
            "对失败或需复算任务执行重跑。",
            "对不再需要执行的任务执行中止或删除。"
        ],
        "任务状态按审计人员操作更新；完成任务的结果可在结果树中查看。",
        ["删除任务可能影响关联结果或子任务，生产系统如有确认弹窗，应先核对影响范围。"],
        ["截图-19"],
    )
    add_operation(
        doc,
        "7.4 处理跑批子任务",
        "跑批任务中部分子任务失败、排队过久或需要单独复核。",
        "已创建跑批任务，任务列表中可进入子任务明细。",
        "在「任务」页签中打开跑批父任务，进入子任务明细。",
        [
            "在任务列表中打开跑批父任务。",
            "进入子任务列表，按成功、失败、排队中或执行中筛选。",
            "打开单个子任务，查看其输入对象、引用资源和失败原因。",
            "对单个失败子任务执行重跑，或在父任务菜单中选择仅失败重跑。",
            "对确认无效的失败子任务执行清理。"
        ],
        "子任务按处理结果更新状态，跑批父任务的整体进度随之刷新。",
        ["重跑前先检查标识列、资源解析状态和技能配置，避免重复执行同一错误配置。"],
        ["截图-19"],
    )


def add_section_8(doc):
    doc.add_heading("8 结果管理", level=1)
    doc.add_paragraph("结果管理用于组织、复核和导出分析产物。结果树中的内容既可以来自任务，也可以来自对话保存或人工编辑。")
    add_operation(
        doc,
        "8.1 查看和组织结果树",
        "需要集中管理当前工作台的分析产物。",
        "当前工作台已有结果或文件夹。",
        "在审计助手右侧切换到「结果」页签，使用结果树。",
        [
            "在右侧切换到结果页签。",
            "通过搜索或排序定位结果。",
            "使用文件夹组织结果层级。",
            "需要批量处理时，进入多选或批量操作状态。",
            "打开目标结果进行预览。"
        ],
        "审计人员可按目录结构找到目标结果，并确认其来源、状态和内容。",
        ["结果目录应按业务主题或任务批次组织，避免全部堆在根目录。"],
        ["截图-20"],
    )
    add_operation(
        doc,
        "8.2 编辑、保存和下载结果",
        "需要对分析结果进行复核、修订或形成可交付材料。",
        "目标结果已生成；当前账号具备编辑或下载权限。",
        "在「结果」页签打开目标结果预览，使用编辑、保存或下载入口。",
        [
            "打开结果预览。",
            "查看基本信息、输出内容和引用来源。",
            "点击编辑，按复核意见修改正文。",
            "保存修改，必要时查看历史版本。",
            "按生产系统支持的格式下载或导出。"
        ],
        "结果内容被保存为新版本，或被下载为后续流转材料。",
        ["导出格式和文件样式以生产系统实际能力为准；下载前应确认结果已复核。"],
        ["截图-21"],
    )
    add_operation(
        doc,
        "8.3 打包下载结果",
        "需要一次性导出多个结果或文件夹。",
        "结果树中已有要打包的结果；当前账号具备下载权限。",
        "在「结果」页签的结果树中选择结果或文件夹，点击打包下载入口。",
        [
            "在结果树中选择一个或多个结果或文件夹。",
            "点击打包下载入口。",
            "确认任务名称、目录结构和输出格式。",
            "提交打包任务。",
            "在任务列表中查看打包进度，完成后下载。"
        ],
        "系统创建文件打包任务，任务完成后可从任务列表下载产物。",
        ["大批量打包可能需要等待；如打包失败，可检查结果是否仍存在或权限是否变化。"],
        ["截图-19", "截图-20"],
    )


def add_section_9(doc):
    doc.add_heading("9 典型审计场景", level=1)
    doc.add_paragraph(
        "本章预留典型审计场景写作框架。具体审计判断口径、资料清单、风险点和结果样式需与业务专家确认后补充；本版不写未确认的专业结论。"
    )
    scenes = [
        ("合同与付款资料比对", "建议补充适用场景、所需资料、比对字段、常见异常和结果复核方式。需业务专家确认合同类型、付款凭证口径、异常判断规则和结果呈现样式。"),
        ("会议纪要事项提取", "建议补充会议纪要资料要求、事项提取范围、责任主体和时间节点。需业务专家确认事项分类、责任认定边界，以及是否形成问题线索。"),
        ("投资项目资料核查", "建议围绕立项、招投标、合同、验收、付款等资料链路组织。需业务专家确认项目阶段划分、关键资料缺失规则和风险提示口径。"),
        ("供应商关联关系分析", "建议补充供应商清单、工商或图谱资源、关系查看和复核路径。需业务专家确认关联关系类型、风险等级表述和是否纳入审计疑点。"),
    ]
    for name, body in scenes:
        add_labeled_paragraph(doc, name, body)

    doc.add_heading("9.1 场景章节建议结构", level=2)
    add_bullets(doc, [
        "场景目标：说明该场景解决的审计业务问题，不写成系统功能介绍。",
        "适用资料：列出建议准备的资料类型，例如合同、付款凭证、验收材料、会议纪要或供应商清单。",
        "操作路径：说明从哪个工作台、资源区、技能或任务入口开始操作。",
        "分析过程：说明需要选择的资源、技能和任务要求，但不写未确认的专业判断规则。",
        "结果复核：说明审计人员如何查看引用资料、结果来源、任务记录和可下载材料。",
        "异常处理：说明资料缺失、解析失败、结果为空或结果不符合预期时的处理建议。",
    ])

    doc.add_heading("9.2 场景输出模板", level=2)
    add_bullets(doc, [
        "场景名称：使用审计业务语言，例如“合同与付款资料比对”，避免写成按钮或页面名称。",
        "业务目标：说明审计人员希望确认什么问题、缩小什么范围或形成什么辅助材料。",
        "前置资料：列出建议准备的资料类型，并说明资料应先完成上传、解析或授权。",
        "推荐操作路径：按实际系统能力说明入口、资源选择、技能使用、任务创建和结果查看路径。",
        "完成标志：说明结果生成、可预览、可下载或可继续复核的可观察状态。",
        "复核要点：说明应回看哪些资料、字段、原文或任务记录；专业判断规则由业务专家确认。",
    ])


def add_section_10(doc):
    doc.add_heading("10 系统管理入口", level=1)
    doc.add_paragraph("系统管理主要面向管理员。对一线审计人员而言，本章只说明入口和与业务使用相关的可见信息，不展开管理员配置流程。")
    add_operation(
        doc,
        "10.1 查看系统管理入口",
        "需要确认系统基础配置入口或联系管理员定位配置项。",
        "当前账号可见系统管理导航。",
        "在顶部导航进入「系统管理」。",
        [
            "在顶部导航进入「系统管理」。",
            "查看左侧菜单，如部门、角色、用户、数据源、模型等。",
            "如需要修改配置，联系具备管理权限的人员处理。"
        ],
        "系统管理页面打开，左侧菜单显示当前账号可见的管理入口。",
        ["不同账号看到的系统管理菜单可能不同。"],
        ["截图-23"],
    )
    add_operation(
        doc,
        "10.2 了解数据源管理入口",
        "业务分析涉及库表或图谱时，需要知道数据源由哪里维护。",
        "当前账号可见数据源管理入口。",
        "在「系统管理」左侧菜单中选择数据源管理入口。",
        [
            "进入系统管理。",
            "选择数据源管理。",
            "查看数据源名称、类型和状态。",
            "如需新增、编辑或测试连接，由管理员按实际权限执行。"
        ],
        "审计人员可定位数据源管理入口，并向管理员准确描述需要调整的数据源。",
        ["本手册不写数据源连接参数、账号或密钥。"],
        ["截图-24"],
    )


def add_section_11(doc):
    doc.add_page_break()
    doc.add_heading("11 FAQ 与状态说明", level=1)
    doc.add_heading("11.1 常见状态", level=2)
    table = doc.add_table(rows=1, cols=4)
    for idx, value in enumerate(["状态", "含义", "审计人员动作", "完成标志"]):
        table.rows[0].cells[idx].text = value
    rows = [
        ("上传中", "文件正在上传到系统。", "保持页面或按系统提示等待；如提供取消入口，可取消不需要的文件。", "状态进入排队中、解析中、完成或失败。"),
        ("排队中", "任务或资料等待系统分配处理资源。", "等待或查看排队提示；无需重复提交。", "状态进入执行中或解析中。"),
        ("解析中", "系统正在识别或处理资料内容。", "等待处理完成；需要时查看进度。", "资料变为完成并可预览或引用。"),
        ("执行中", "任务正在调用技能和资源进行分析。", "在任务列表查看进度，不重复创建同类任务。", "任务进入完成或失败。"),
        ("完成", "资料、任务或结果已处理完成。", "打开结果或资料进行复核。", "可预览、引用、编辑或下载。"),
        ("失败", "上传、解析或任务执行未成功。", "查看失败原因，必要时重跑、重新上传或调整资源。", "重跑后进入完成，或明确不再处理。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, widths=[1.0, 2.0, 2.0, 1.3])

    doc.add_heading("11.2 常见问题", level=2)
    faq = [
        ("看不到某个工作台怎么办？", "先切换「我的工作台 / 共享工作台」，清空搜索条件；仍不可见时联系管理员确认权限。"),
        ("资料一直处于排队或解析中怎么办？", "先确认是否为大文件或批量资料；如长时间无变化，记录资料名称和时间点后反馈运维或管理员。"),
        ("分析结果不符合预期怎么办？", "先检查引用资料、技能和任务要求是否准确；必要时补充资料、修改任务要求后重跑。"),
        ("什么时候应该使用跑批任务？", "当同一分析规则需要对多行对象或多个主体重复执行时使用；单个问题优先使用单次任务或对话追问。"),
        ("对话结果能直接作为正式结论吗？", "不能直接作为最终结论。应保存为结果后由审计人员复核、编辑和确认。"),
        ("技能库和工作台技能有什么区别？", "技能库是可复用资产的管理位置；工作台技能是当前工作台已引用或创建、可直接用于本次审计任务的技能。"),
        ("导出格式与手册不一致怎么办？", "以生产系统实际支持格式为准；手册不写死未确认的导出能力。"),
    ]
    for q, a in faq:
        add_labeled_paragraph(doc, q, a)


def add_appendices(doc):
    doc.add_heading("附录 待确认事项", level=1)
    rows = [
        ("生产系统访问地址与登录方式", "不同部署环境可能不同。", "正文只写“进入系统”，不写具体地址。"),
        ("文件格式、大小、数量限制", "部署配置可能随客户环境变化。", "仅提示以生产配置为准。"),
        ("权限角色与共享规则", "真实权限受组织、角色、租户配置影响。", "只写可见操作和联系管理员。"),
        ("导出格式和下载样式", "生产环境可能分阶段开放。", "写为“按系统支持格式下载”。"),
        ("知识库资源能力", "该能力是否开放取决于部署环境。", "正文说明入口，不展开正式流程。"),
        ("系统管理配置项", "一线审计人员通常不直接维护。", "仅做入口说明，不写管理员手册。"),
    ]
    for item, reason, handling in rows:
        add_labeled_paragraph(doc, item, f"为什么需要确认：{reason} 手册口径：{handling}")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_document_notes_page(doc)
    add_manual_toc(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)
    add_section_10(doc)
    add_section_11(doc)
    add_appendices(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
