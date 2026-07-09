from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/mac/Desktop/workspace")
SRC = ROOT / "x 工作日志/2026年上半年年中总结报告.md"
OUT = ROOT / "x 工作日志/2026年上半年年中总结报告.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)


def set_east_asia(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name="Microsoft YaHei", size=11, color=None, bold=False):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = color
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_bottom_border(paragraph, color="D9E2F3", size="6"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_runs(paragraph, text, base_bold=False, color=None, size=None):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = base_bold
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            bold = True
        run = paragraph.add_run(part)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
        set_east_asia(run)


def set_para(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Bullet 2", "List Bullet 3"]:
        if style_name in styles:
            set_style_font(styles[style_name], size=11)
            styles[style_name].paragraph_format.space_after = Pt(5)
            styles[style_name].paragraph_format.line_spacing = 1.1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("2026年上半年年中总结报告")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    set_east_asia(run)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    in_refs = False
    first_title = True
    previous_blank = False

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            previous_blank = True
            continue
        stripped = line.strip()

        if stripped == "---":
            p = doc.add_paragraph()
            set_para(p, before=6, after=8)
            add_bottom_border(p)
            continue

        if stripped == "信息引用":
            in_refs = True
            p = doc.add_paragraph(style="Heading 2")
            add_runs(p, stripped)
            continue

        if in_refs:
            p = doc.add_paragraph()
            set_para(p, before=0, after=2, line=1.0)
            add_runs(p, stripped, color=MUTED, size=9)
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if first_title:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para(p, before=0, after=14, line=1.15)
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(22)
                r.font.color.rgb = RGBColor(31, 58, 95)
                set_east_asia(r)
                first_title = False
            else:
                p = doc.add_paragraph(style="Heading 1")
                add_runs(p, text)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            style = "Heading 1" if level <= 1 else "Heading 2"
            p = doc.add_paragraph(style=style)
            add_runs(p, text)
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_runs(p, stripped[3:].strip())
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.*)$", raw)
        if bullet_match:
            spaces, text = bullet_match.groups()
            depth = 0 if len(spaces) <= 1 else 1 if len(spaces) <= 4 else 2
            style = ["List Bullet", "List Bullet 2", "List Bullet 3"][depth]
            p = doc.add_paragraph(style=style)
            add_runs(p, text)
            continue

        p = doc.add_paragraph()
        set_para(p, before=0, after=8 if previous_blank else 6)
        add_runs(p, stripped)
        previous_blank = False

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(make_doc())
